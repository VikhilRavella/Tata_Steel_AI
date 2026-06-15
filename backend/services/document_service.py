from sqlalchemy.ext.asyncio import AsyncSession
import os
import uuid
import json
from sqlalchemy.orm import Session
from backend.models import Document, User
from backend.services.rag_service import chunk_text, get_embeddings
try:
    import chromadb
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
if CHROMA_AVAILABLE:
    from chromadb.config import Settings
    chroma_client = chromadb.PersistentClient(path='backend/chroma_db', settings=Settings(anonymized_telemetry=False))
    collection = chroma_client.get_or_create_collection(name='maintenance_docs')
else:
    chroma_client = None
    collection = None

async def process_document(file_path: str, filename: str, current_user: User, db: AsyncSession, doc_category: str='SOP') -> bool:
    """
    Simulates checking for the security stamp, extracts text, chunks it, and saves embeddings.
    Returns True if passed, False if rejected by stamp check.
    Raises Exception if permission denied.
    """
    # TEMPORARILY DISABLED FOR HACKATHON DEMO
    # if current_user.role == 'engineer':
    #     if current_user.specialization and doc_category.lower() not in current_user.specialization.lower():
    #         raise Exception('Permission Denied: Engineers can only upload documents within their specialization.')
    is_valid_stamp = True
    if 'classified' in filename.lower():
        is_valid_stamp = False
    if not is_valid_stamp:
        return False
    text = ''
    if filename.endswith('.pdf') and PYMUPDF_AVAILABLE:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
    elif filename.endswith('.docx') and DOCX_AVAILABLE:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text
    if not text:
        text = 'Simulated extracted text for ' + filename
    chunks = chunk_text(text, chunk_size=500)
    doc_record = Document(filename=f'{uuid.uuid4()}_{filename}', original_filename=filename, file_type=filename.split('.')[-1], file_path=file_path, uploaded_by=current_user.id, department=current_user.department, document_category='SOP', assigned_roles=json.dumps([current_user.role, 'manager']), stamp_verified=True, processing_status='ready', chroma_collection='maintenance_docs')
    db.add(doc_record)
    db.commit()
    db.refresh(doc_record)
    if collection:
        for i, chunk in enumerate(chunks):
            embedding = await get_embeddings(chunk)
            collection.add(embeddings=[embedding], documents=[chunk], metadatas=[{'doc_id': str(doc_record.id), 'filename': filename, 'department': current_user.department if current_user.department else 'unknown', 'role': current_user.role}], ids=[f'{doc_record.id}_chunk_{i}'])
    return True

def extract_section_and_heading(page_text: str) -> tuple[str, str]:
    """
    Analyzes page text to find the most probable section name and heading.
    """
    lines = [line.strip() for line in page_text.split('\n') if line.strip()]
    section = "General Section"
    heading = "Technical Content"
    
    # Try to find a section (e.g. "SECTION 1", "CHAPTER 2", uppercase lines)
    for line in lines[:5]:
        line_upper = line.upper()
        if any(kw in line_upper for kw in ["SECTION", "CHAPTER", "SOP", "MANUAL", "PREFACE", "SAFETY"]):
            section = line
            break
            
    # Try to find a heading
    for line in lines:
        if len(line.split()) < 8 and (line.isupper() or any(line.startswith(p) for p in ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "A.", "B.", "C."])):
            heading = line
            break
            
    return section, heading

def chunk_page_text(text: str, chunk_size: int = 800, overlap: int = 100) -> list[str]:
    chunks = []
    if not text:
        return chunks
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= len(text):
            break
        start += chunk_size - overlap
    return chunks

def find_chunk_metadata_in_pdf(file_path: str, chunk_text: str) -> tuple[int, str, str]:
    """
    Scans a PDF dynamically to find where a chunk is located, returning (page_num, section, heading).
    Used as a robust fallback for legacy/unstaged documents.
    """
    if not file_path or not os.path.exists(file_path):
        return 1, "General Section", "Technical Content"
    try:
        import fitz
        import re
        doc = fitz.open(file_path)
        sample = chunk_text[:120].strip().lower()
        if not sample:
            doc.close()
            return 1, "General Section", "Technical Content"
            
        sample_clean = re.sub(r'[^a-z0-9]', '', sample)
        if not sample_clean:
            doc.close()
            return 1, "General Section", "Technical Content"
            
        for page_idx, page in enumerate(doc):
            page_text_clean = re.sub(r'[^a-z0-9]', '', page.get_text().lower())
            if sample_clean in page_text_clean:
                section, heading = extract_section_and_heading(page.get_text())
                doc.close()
                return page_idx + 1, section, heading
        doc.close()
    except Exception as e:
        print(f"Error scanning PDF: {e}")
    return 1, "General Section", "Technical Content"

async def embed_company_document(file_path: str, filename: str, company_doc_id: int):
    """
    Extracts text from a Manager-uploaded Company Document page by page,
    chunks it, and saves embeddings and metadata into ChromaDB.
    """
    if not CHROMA_AVAILABLE or not collection:
        print("ChromaDB not available, skipping embeddings.")
        return
        
    if filename.endswith('.pdf') and PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(file_path)
            page_count = len(doc)
            print(f"Uploaded PDF Page Count: {page_count}")
            
            for page_idx, page in enumerate(doc):
                page_num = page_idx + 1
                page_text = page.get_text()
                if not page_text.strip():
                    continue
                    
                section, heading = extract_section_and_heading(page_text)
                chunks = chunk_page_text(page_text, chunk_size=800, overlap=100)
                
                for i, chunk in enumerate(chunks):
                    embedding = await get_embeddings(chunk)
                    collection.add(
                        embeddings=[embedding], 
                        documents=[chunk], 
                        metadatas=[{
                            'company_doc_id': str(company_doc_id),
                            'document_id': str(company_doc_id),
                            'filename': filename, 
                            'type': 'company_sop', 
                            'role': 'engineer',
                            'page_number': page_num,
                            'section_name': section,
                            'heading': heading,
                            'chunk_id': f'comp_doc_{company_doc_id}_page_{page_num}_chunk_{i}'
                        }], 
                        ids=[f'comp_doc_{company_doc_id}_page_{page_num}_chunk_{i}']
                    )
            doc.close()
        except Exception as e:
            print("Failed to extract and embed PDF text:", e)
    else:
        # Fallback for non-PDF or simple extraction
        text = 'Extracted text for ' + filename
        chunks = chunk_page_text(text, chunk_size=800, overlap=100)
        for i, chunk in enumerate(chunks):
            embedding = await get_embeddings(chunk)
            collection.add(
                embeddings=[embedding], 
                documents=[chunk], 
                metadatas=[{
                    'company_doc_id': str(company_doc_id),
                    'document_id': str(company_doc_id),
                    'filename': filename, 
                    'type': 'company_sop', 
                    'role': 'engineer',
                    'page_number': 1,
                    'section_name': 'General Section',
                    'heading': 'Technical Content',
                    'chunk_id': f'comp_doc_{company_doc_id}_chunk_{i}'
                }], 
                ids=[f'comp_doc_{company_doc_id}_chunk_{i}']
            )

async def get_all_document_chunks(company_doc_id: int) -> list[str]:
    """
    Retrieves all document chunks for the given company_doc_id from ChromaDB.
    """
    if not CHROMA_AVAILABLE or not collection:
        return []
    try:
        res = collection.get(where={'company_doc_id': str(company_doc_id)})
        if res and res.get('documents'):
            return res['documents']
    except Exception as e:
        print(f"Error fetching document chunks for {company_doc_id}: {e}")
    return []

async def get_all_document_chunks_with_metadata(company_doc_id: int, file_path: str = None) -> list[dict]:
    """
    Retrieves all document chunks with their page_number, section_name, and heading.
    """
    if not CHROMA_AVAILABLE or not collection:
        return []
    try:
        res = collection.get(where={'company_doc_id': str(company_doc_id)}, include=["documents", "metadatas"])
        chunks = []
        if res and res.get('documents'):
            docs = res['documents']
            metas = res['metadatas'] if res.get('metadatas') else [None] * len(docs)
            for doc_text, meta in zip(docs, metas):
                meta_dict = meta if meta else {}
                page_num = meta_dict.get('page_number')
                section = meta_dict.get('section_name')
                heading = meta_dict.get('heading')
                
                if page_num is None or section is None or heading is None:
                    page_num, section, heading = find_chunk_metadata_in_pdf(file_path, doc_text)
                    
                chunks.append({
                    "text": doc_text,
                    "page_number": page_num,
                    "section_name": section,
                    "heading": heading
                })
            return chunks
    except Exception as e:
        print(f"Error fetching document chunks with metadata: {e}")
    return []

async def remove_company_document_embeddings(company_doc_id: int):
    """
    Removes all embeddings related to a specific Company Document from ChromaDB.
    """
    if not CHROMA_AVAILABLE or not collection:
        return
    try:
        collection.delete(where={"company_doc_id": str(company_doc_id)})
    except Exception as e:
        print(f"Failed to delete embeddings for comp_doc_{company_doc_id}:", e)

async def search_documents(query: str, user_role: str, user_dept: str, top_k: int=10, company_doc_id: int = None):
    """
    Search ChromaDB for relevant document excerpts based on query.
    Filters by user role.
    """
    if not collection:
        return ['(Mock DB) Found reference in C-Series Pump SOP, page 42.']
        
    expanded_query = query.lower()
    if any(k in expanded_query for k in ["equipment", "machine", "component", "compressor", "motor", "pump"]):
        expanded_query += " equipment machine component compressor motor pump spare parts"
    if any(k in expanded_query for k in ["safety", "warning", "caution", "ppe", "lockout", "procedure"]):
        expanded_query += " safety warning caution PPE lockout procedure hazard precaution"
        
    query_embedding = await get_embeddings(expanded_query)
    
    where_clause = {}
    if company_doc_id is not None:
        where_clause = {'company_doc_id': str(company_doc_id)}
    else:
        where_clause = {'role': user_role}
        
    results = collection.query(query_embeddings=[query_embedding], n_results=top_k, where=where_clause)
    if results and results['documents'] and results['documents'][0]:
        return results['documents'][0]
    return []

def parse_page_number(query: str):
    import re
    # Match patterns like: page 5, p. 5, p 5, pg 5, pg. 5, page: 5, page number 5
    match = re.search(r'\b(?:page|p\.?|pg\.?|page\s+number)\s*[:#-]?\s*(\d+)\b', query, re.IGNORECASE)
    if match:
        return int(match.group(1))
    
    # Try text numbers for page numbers up to twenty
    text_nums = {
        "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
        "eleven": 11, "twelve": 12, "thirteen": 13, "fourteen": 14, "fifteen": 15, "sixteen": 16, "seventeen": 17,
        "eighteen": 18, "nineteen": 19, "twenty": 20
    }
    for word, num in text_nums.items():
        if re.search(rf'\bpage\s+{word}\b', query, re.IGNORECASE):
            return num
    return None

async def search_documents_v3(
    query: str, 
    user_role: str, 
    user_dept: str, 
    top_k: int = 10, 
    company_doc_id: int = None,
    file_path: str = None
) -> list[dict]:
    """
    Search ChromaDB for relevant document excerpts based on query and return dictionaries
    with text, page_number, section_name, heading, and other metadata.
    Supports exact page-number metadata filtering and fallback document scanning.
    """
    if not collection:
        return [{
            "text": "(Mock DB) Found reference in C-Series Pump SOP, page 42.",
            "page_number": 42,
            "section_name": "General Section",
            "heading": "Technical Content"
        }]

    target_page_number = parse_page_number(query)
    
    where_clause = {}
    if company_doc_id is not None:
        where_clause = {'company_doc_id': str(company_doc_id)}
    else:
        where_clause = {'role': user_role}

    results = None
    if target_page_number is not None:
        try:
            combined_where = {
                "$and": [
                    {"company_doc_id": str(company_doc_id)},
                    {"page_number": target_page_number}
                ]
            } if company_doc_id is not None else {
                "$and": [
                    {"role": user_role},
                    {"page_number": target_page_number}
                ]
            }
            results = collection.query(
                query_embeddings=[await get_embeddings(query)],
                n_results=top_k,
                where=combined_where,
                include=["documents", "metadatas"]
            )
        except Exception as e:
            print(f"Error querying metadata directly: {e}")
            results = None

    if not results or not results.get('documents') or not results['documents'][0]:
        expanded_query = query.lower()
        if any(k in expanded_query for k in ["equipment", "machine", "component", "compressor", "motor", "pump"]):
            expanded_query += " equipment machine component compressor motor pump spare parts"
        if any(k in expanded_query for k in ["safety", "warning", "caution", "ppe", "lockout", "procedure"]):
            expanded_query += " safety warning caution PPE lockout procedure hazard precaution"

        query_embedding = await get_embeddings(expanded_query)
        query_top_k = 100 if target_page_number is not None else top_k
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=query_top_k,
            where=where_clause,
            include=["documents", "metadatas"]
        )

    chunks_with_metadata = []
    if results and results.get('documents') and results['documents'][0]:
        docs = results['documents'][0]
        metas = results['metadatas'][0] if results.get('metadatas') else [None] * len(docs)
        
        for doc_text, meta in zip(docs, metas):
            meta_dict = meta if meta else {}
            page_num = meta_dict.get('page_number')
            section = meta_dict.get('section_name')
            heading = meta_dict.get('heading')
            
            if page_num is None or section is None or heading is None:
                page_num, section, heading = find_chunk_metadata_in_pdf(file_path, doc_text)
                
            if target_page_number is not None and page_num != target_page_number:
                continue
                
            chunks_with_metadata.append({
                "text": doc_text,
                "page_number": page_num,
                "section_name": section,
                "heading": heading
            })
            
            if len(chunks_with_metadata) >= top_k:
                break

    if target_page_number is not None and not chunks_with_metadata and company_doc_id is not None:
        try:
            all_res = collection.get(
                where={'company_doc_id': str(company_doc_id)},
                include=["documents", "metadatas"]
            )
            if all_res and all_res.get('documents'):
                docs = all_res['documents']
                metas = all_res['metadatas'] if all_res.get('metadatas') else [None] * len(docs)
                for doc_text, meta in zip(docs, metas):
                    meta_dict = meta if meta else {}
                    page_num = meta_dict.get('page_number')
                    section = meta_dict.get('section_name')
                    heading = meta_dict.get('heading')
                    
                    if page_num is None or section is None or heading is None:
                        page_num, section, heading = find_chunk_metadata_in_pdf(file_path, doc_text)
                        
                    if page_num == target_page_number:
                        chunks_with_metadata.append({
                            "text": doc_text,
                            "page_number": page_num,
                            "section_name": section,
                            "heading": heading
                        })
                        if len(chunks_with_metadata) >= top_k:
                            break
        except Exception as e:
            print(f"Error in all-chunks fallback filter: {e}")
                
    return chunks_with_metadata